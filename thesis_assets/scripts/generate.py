#!/usr/bin/env python3
"""论文生成脚本 — 微服务核心组件实验报告"""
import os, copy
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PROJECT = os.path.dirname(BASE)
TEMPLATE_FILE = os.path.join(PROJECT, 'docx', '学号_姓名_《微服务核心组件实验》实验报告.docx')
OUTPUT_FILE = os.path.join(BASE, 'output', '202502150280_吴梓沐_《微服务核心组件实验》实验报告.docx')
IMAGES_DIR = os.path.join(BASE, 'images')

# ── User Info ────────────────────────────────────────────
STUDENT_ID = '202502150280'
NAME = '吴梓沐'
CLASS_NAME = 'a25计算机科学3班'
DATE = '2026年5月31日'

# ── Experiment Environment ───────────────────────────────
ENV_INFO = {
    '操作系统': 'macOS 26 Sequoia',
    'JDK版本': 'JDK 17.0.14 (Oracle OpenJDK)',
    'Maven版本': 'Apache Maven 3.9.6',
    'IDE版本': 'IntelliJ IDEA 2025.1 (Ultimate Edition)',
}


# ── Helpers ──────────────────────────────────────────────

def replace_paragraph(p, new_text, font_size=None, bold=False):
    """Clear paragraph and set new text, preserving paragraph style."""
    saved_style = p.style
    p.clear()
    run = p.add_run(new_text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    elif saved_style and 'Heading' in str(saved_style.name):
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return run


def clear_paragraph(p):
    """Clear paragraph content."""
    p.clear()


def add_image_after_paragraph(doc, anchor_para, img_filename, caption_text, width_inches=5.0):
    """Insert image + caption after the given paragraph.

    Uses python-docx's built-in API: creates paragraphs at the end of the document
    then moves them right after the anchor paragraph in the XML body.
    """
    img_path = os.path.join(IMAGES_DIR, img_filename)
    if not os.path.exists(img_path):
        print(f'  WARNING: Image not found: {img_path}')
        return

    body = doc.element.body
    anchor_elem = anchor_para._element
    anchor_idx = list(body).index(anchor_elem)

    def make_centered_paragraph():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    # Spacer before
    spacer_before = doc.add_paragraph()

    # Image paragraph
    img_para = make_centered_paragraph()
    run = img_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))

    # Caption paragraph
    cap_para = make_centered_paragraph()
    cap_run = cap_para.add_run(caption_text)
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = None  # default color

    # Spacer after
    spacer_after = doc.add_paragraph()

    # Move elements: we need to move spacer_before, img_para, cap_para, spacer_after
    # to right after anchor_elem
    new_elems = [spacer_before._element, img_para._element, cap_para._element, spacer_after._element]

    # anchor_idx is the position of the anchor in the body
    # Insert all new elements after it (in reverse order to maintain positions)
    insert_at = anchor_idx + 1
    for elem in reversed(new_elems):
        body.remove(elem)
        body.insert(insert_at, elem)

    print(f'  Inserted image: {caption_text}')


# ── Fill functions ───────────────────────────────────────

def fill_cover_table(doc):
    """Fill the cover page table with user info."""
    table = doc.tables[0]
    # Row 0: 课程名称
    table.rows[0].cells[1].paragraphs[0].clear()
    table.rows[0].cells[1].paragraphs[0].add_run('微服务应用开发')
    # Row 1: 实验名称
    table.rows[1].cells[1].paragraphs[0].clear()
    table.rows[1].cells[1].paragraphs[0].add_run('微服务核心组件实验')
    # Row 2: 学号
    table.rows[2].cells[1].paragraphs[0].clear()
    table.rows[2].cells[1].paragraphs[0].add_run(STUDENT_ID)
    # Row 3: 姓名
    table.rows[3].cells[1].paragraphs[0].clear()
    table.rows[3].cells[1].paragraphs[0].add_run(NAME)
    # Row 4: 班级
    table.rows[4].cells[1].paragraphs[0].clear()
    table.rows[4].cells[1].paragraphs[0].add_run(CLASS_NAME)
    # Row 6: 实验日期
    table.rows[6].cells[1].paragraphs[0].clear()
    table.rows[6].cells[1].paragraphs[0].add_run(DATE)
    # Row 7: 报告日期
    table.rows[7].cells[1].paragraphs[0].clear()
    table.rows[7].cells[1].paragraphs[0].add_run(DATE)
    print('  Cover table filled.')


def fill_experiment_env(doc):
    """Fill experiment environment info."""
    replace_paragraph(doc.paragraphs[15], f'操作系统：{ENV_INFO["操作系统"]}')
    replace_paragraph(doc.paragraphs[16], f'JDK版本：{ENV_INFO["JDK版本"]}')
    replace_paragraph(doc.paragraphs[17], f'Maven版本：{ENV_INFO["Maven版本"]}')
    replace_paragraph(doc.paragraphs[18], f'IDE版本：{ENV_INFO["IDE版本"]}')
    print('  Experiment environment filled.')


def replace_postman_refs(doc):
    """Replace 'Postman' references with 'curl'."""
    count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if 'Postman' in run.text:
                run.text = run.text.replace('Postman', 'curl')
                count += 1
    print(f'  Replaced Postman references in {count} runs.')


def insert_all_screenshots(doc):
    """Insert screenshots at designated positions.

    IMPORTANT: Screenshots are inserted sequentially. After the first image insertion
    at paragraph 35, the paragraph count increases. We need to work backwards or
    insert from the end forward to keep indices stable.
    """

    # Work from END to BEGINNING so indices don't shift under us.

    # --- Slot 4: after paragraph [96] (Sentinel: 4 images) ---
    # Insert in REVERSE order (9,8,7,6) so they appear in natural order (6,7,8,9)
    clear_paragraph(doc.paragraphs[96])
    anchor = doc.paragraphs[96]
    add_image_after_paragraph(doc, anchor, '09_gateway_limit.png',
        '图9 网关层 Sentinel 限流超出阈值响应', width_inches=4.8)
    add_image_after_paragraph(doc, anchor, '08_circuit_breaker.png',
        '图8 熔断器打开后 Feign Fallback 降级返回', width_inches=4.8)
    add_image_after_paragraph(doc, anchor, '07_sentinel_blocked.png',
        '图7 Sentinel 限流拦截响应', width_inches=4.8)
    add_image_after_paragraph(doc, anchor, '06_sentinel_dashboard.png',
        '图6 Sentinel Dashboard 流控与熔断规则配置', width_inches=4.8)

    # --- Slot 3: after paragraph [71] (Gateway+JWT: 3 images) ---
    # Insert in REVERSE order (5,4,3) so they appear in natural order (3,4,5)
    clear_paragraph(doc.paragraphs[71])
    anchor = doc.paragraphs[71]
    add_image_after_paragraph(doc, anchor, '05_401_error.png',
        '图5 无 Token 请求返回 401', width_inches=4.8)
    add_image_after_paragraph(doc, anchor, '04_token_success.png',
        '图4 携带 Token 调用业务接口成功', width_inches=4.8)
    add_image_after_paragraph(doc, anchor, '03_login_response.png',
        '图3 登录接口返回 Token', width_inches=4.8)

    # --- Slot 2: after paragraph [48] (API call: 1 image) ---
    clear_paragraph(doc.paragraphs[48])
    add_image_after_paragraph(doc, doc.paragraphs[48], '02_api_response.png',
        '图2 API 远程调用测试结果', width_inches=4.8)

    # --- Slot 1: after paragraphs [34-35] (Nacos: 1 image) ---
    clear_paragraph(doc.paragraphs[34])
    clear_paragraph(doc.paragraphs[35])
    add_image_after_paragraph(doc, doc.paragraphs[35], '01_nacos_console.png',
        '图1 Nacos 控制台服务列表', width_inches=4.8)

    print('  All screenshots inserted.')


# ── Main ─────────────────────────────────────────────────

def main():
    print(f'Template: {TEMPLATE_FILE}')
    print(f'Output:   {OUTPUT_FILE}')

    doc = Document(TEMPLATE_FILE)

    print('\n[1/4] Filling cover table...')
    fill_cover_table(doc)

    print('\n[2/4] Filling experiment environment...')
    fill_experiment_env(doc)

    print('\n[3/4] Replacing Postman references...')
    replace_postman_refs(doc)

    print('\n[4/4] Inserting screenshots...')
    insert_all_screenshots(doc)

    # Save
    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f'\nDone! Output saved to: {OUTPUT_FILE}')


if __name__ == '__main__':
    main()
